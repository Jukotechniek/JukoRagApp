# Document processing API endpoint
# Processes uploaded documents for RAG: extracts text, chunks, and generates embeddings

import os
from dotenv import load_dotenv
import uuid
import time
from typing import Optional

# import langchain
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.document_loaders import CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document

# For DOCX files, use python-docx directly
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX files will not be supported.")

# import supabase
from supabase.client import Client, create_client

# import FastAPI
from fastapi import FastAPI, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# load environment variables
load_dotenv()

# initiating supabase
supabase_url = os.environ.get("SUPABASE_URL")
supabase_key = os.environ.get("SUPABASE_SERVICE_KEY")
supabase: Client = create_client(supabase_url, supabase_key)

# initiating embeddings model
embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

# FastAPI app - can be merged with chat.py later
app = FastAPI(title="Document Processing API")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Request/Response models
class ProcessDocumentRequest(BaseModel):
    documentId: str
    organizationId: str

class ProcessDocumentResponse(BaseModel):
    success: bool
    message: str
    chunksProcessed: Optional[int] = None
    error: Optional[str] = None


def extract_text_from_file(file_path: str, file_type: str, file_name: str) -> str:
    """Extract text from various file types"""
    try:
        if file_type == 'application/pdf' or file_name.lower().endswith('.pdf'):
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_name.lower().endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not installed. Install it with: pip install python-docx")
            
            doc = DocxDocument(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            return "\n\n".join(paragraphs)
        
        elif file_type == 'text/plain' or file_name.lower().endswith('.txt'):
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        
        elif file_type == 'text/csv' or file_name.lower().endswith('.csv'):
            loader = CSVLoader(file_path, encoding='utf-8')
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        
        else:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    
    except Exception as e:
        raise Exception(f"Failed to extract text from {file_name}: {str(e)}")


@app.post("/api/process-document", response_model=ProcessDocumentResponse)
async def process_document(
    request: ProcessDocumentRequest,
    authorization: Optional[str] = Header(None)
):
    """Process a document for RAG: extract text, chunk, and generate embeddings"""
    start_time = time.time()
    
    try:
        # Get document info from database
        doc_result = supabase.table("documents").select("file_url, name, file_type").eq("id", request.documentId).single().execute()
        
        if not doc_result.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc = doc_result.data
        
        # Extract path from file_url
        file_url = doc.get("file_url")
        if not file_url:
            raise HTTPException(status_code=400, detail="Document has no file URL")
        
        # Extract storage path from URL
        # URL format: https://...supabase.co/storage/v1/object/documents/path/to/file.pdf
        import re
        url_match = re.search(r'/documents/(.+)$', file_url)
        if not url_match:
            raise HTTPException(status_code=400, detail="Could not extract file path from URL")
        
        storage_path = url_match.group(1)
        
        # Download file from Supabase Storage to temporary location
        import tempfile
        temp_dir = tempfile.mkdtemp()
        file_name = doc.get("name", "document")
        temp_file_path = os.path.join(temp_dir, file_name)
        
        try:
            # Download file from Supabase Storage
            file_response = supabase.storage.from("documents").download(storage_path)
            
            if not file_response:
                raise HTTPException(status_code=500, detail="Failed to download file: No response")
            
            # Supabase Python client returns bytes directly
            file_data = file_response
            
            # Save to temporary file
            with open(temp_file_path, 'wb') as f:
                f.write(file_data)
            
            # Extract text from file
            file_type = doc.get("file_type", "")
            text_content = extract_text_from_file(temp_file_path, file_type, file_name)
            
            if not text_content or len(text_content.strip()) == 0:
                raise HTTPException(status_code=400, detail="No text content extracted from document")
            
            # Split into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1500,
                chunk_overlap=200
            )
            
            # Create LangChain documents with metadata
            langchain_docs = [Document(
                page_content=text_content,
                metadata={
                    "document_id": request.documentId,
                    "source": file_name,
                    "file_type": file_type
                }
            )]
            
            chunks = text_splitter.split_documents(langchain_docs)
            
            # Add document_id to each chunk's metadata
            for chunk in chunks:
                chunk.metadata["document_id"] = request.documentId
            
            print(f"Created {len(chunks)} chunks from document {file_name}")
            
            # Generate embeddings for all chunks
            print(f"Generating embeddings for {len(chunks)} chunks...")
            texts = [chunk.page_content for chunk in chunks]
            embeddings_list = embeddings.embed_documents(texts)
            
            # Delete existing document sections for this document
            delete_result = supabase.table("document_sections").delete().eq("document_id", request.documentId).execute()
            
            if hasattr(delete_result, 'error') and delete_result.error:
                print(f"Warning: Could not delete existing sections: {delete_result.error}")
            
            # Insert document sections with embeddings
            print(f"Inserting {len(chunks)} document sections into database...")
            sections_to_insert = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings_list)):
                sections_to_insert.append({
                    "document_id": request.documentId,
                    "content": chunk.page_content,
                    "metadata": chunk.metadata,
                    "embedding": embedding
                })
                
                # Insert in batches of 10 for better performance
                if len(sections_to_insert) >= 10 or i == len(chunks) - 1:
                    insert_result = supabase.table("document_sections").insert(sections_to_insert).execute()
                    
                    if hasattr(insert_result, 'error') and insert_result.error:
                        raise HTTPException(status_code=500, detail=f"Failed to insert document sections: {insert_result.error}")
                    
                    sections_to_insert = []
                    if (i + 1) % 10 == 0:
                        print(f"Inserted {i + 1}/{len(chunks)} chunks...")
            
            duration = time.time() - start_time
            
            print(f"Successfully processed document {file_name}: {len(chunks)} chunks in {duration:.2f}s")
            
            return ProcessDocumentResponse(
                success=True,
                message=f"Successfully processed document: {len(chunks)} chunks created",
                chunksProcessed=len(chunks)
            )
        
        finally:
            # Clean up temporary file
            try:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                os.rmdir(temp_dir)
            except Exception as cleanup_error:
                print(f"Warning: Could not clean up temporary files: {cleanup_error}")
    
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        print(f"Error processing document: {error_msg}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process document: {error_msg}"
        )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)


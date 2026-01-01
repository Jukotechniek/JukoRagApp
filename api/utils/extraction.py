"""Document extraction utilities for various file types"""
import os
import html
import re
from typing import List
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader

# Check for optional dependencies
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX files will not be supported.")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF (fitz) not available. Footer extraction will be limited.")

try:
    from mistralai import Mistral
    MISTRAL_AVAILABLE = True
except ImportError:
    MISTRAL_AVAILABLE = False
    print("Warning: mistralai not available. Mistral OCR will not be available.")

from .text_utils import normalize_spaced_text, encode_file_to_base64


def extract_footer_info_from_pdf(pdf_path: str, page_num: int) -> dict:
    """Extract footer information from PDF page using area-based extraction with get_text("dict").
    
    Uses structured text extraction with bounding boxes for better accuracy.
    Works for any PDF schema type without schema-specific patterns.
    """
    footer_info = {
        "page_number_footer": None,
        "project_description": None
    }
    
    if not PYMUPDF_AVAILABLE:
        return footer_info
    
    try:
        doc = fitz.open(pdf_path)
        if page_num < len(doc):
            page = doc[page_num]
            
            # Get page dimensions
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            
            # ROI = bottom-right area (rechtsonder) for page number
            # Use same fractions as working code: 85% width, 90% height
            X0_FRAC = 0.85  # vanaf 85% breedte
            Y0_FRAC = 0.90  # vanaf 90% hoogte (laatste regels)
            
            # Use get_text("dict") for better structure (like working code)
            text_dict = page.get_text("dict")
            
            # Extract page number from right bottom region
            page_number_candidates = []
            
            # Search for lines in the bottom-right area
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_bbox = line["bbox"]  # [x0, y0, x1, y1]
                        # Check if line is in ROI (right bottom area)
                        if line_bbox[3] >= page_height * Y0_FRAC and line_bbox[0] >= page_width * X0_FRAC:
                            # Construct text from line (combine all spans)
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            
                            # Normalize text (remove spaces between characters)
                            normalized = normalize_spaced_text(line_text)
                            # Also remove all spaces for number extraction
                            normalized_no_spaces = normalized.replace(' ', '')
                            
                            # Extract all numbers from the normalized text
                            current_number = ""
                            for char in normalized_no_spaces:
                                if char.isdigit():
                                    current_number += char
                                else:
                                    if current_number:
                                        try:
                                            num = int(current_number)
                                            if num >= 1:  # Valid page number
                                                # Store: (number, x_right, y_bottom) for sorting
                                                page_number_candidates.append((num, line_bbox[2], line_bbox[3]))
                                        except ValueError:
                                            pass
                                        current_number = ""
                            # Don't forget the last number if text ends with digits
                            if current_number:
                                try:
                                    num = int(current_number)
                                    if num >= 1:
                                        page_number_candidates.append((num, line_bbox[2], line_bbox[3]))
                                except ValueError:
                                    pass
            
            # Select the best page number candidate
            if page_number_candidates:
                # Sort by y-position (bottom first), then by x-position (right first)
                # This gives us the most bottom-right number
                page_number_candidates.sort(key=lambda x: (-x[2], -x[1]))
                
                # Take the most bottom-right number
                # Prefer larger numbers (>= 100) if available, otherwise use the most bottom-right
                large_candidates = [(n, x, y) for n, x, y in page_number_candidates if n >= 100]
                if large_candidates:
                    # Sort large candidates the same way
                    large_candidates.sort(key=lambda x: (-x[2], -x[1]))
                    footer_info["page_number_footer"] = large_candidates[0][0]
                else:
                    # Use the most bottom-right number even if small
                    footer_info["page_number_footer"] = page_number_candidates[0][0]
            
            # Extract project description and text above from bottom area (bottom 15%)
            # Use get_text("dict") for structured extraction like page number
            all_lines = []
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_bbox = line["bbox"]
                        # Collect lines in bottom 15% of page
                        if line_bbox[3] >= page_height * 0.85:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            if line_text.strip():
                                normalized_text = normalize_spaced_text(line_text)
                                all_lines.append((normalized_text, line_bbox[0], line_bbox[1], line_bbox[2], line_bbox[3]))
            
            if all_lines:
                # Sort by y-position (top to bottom)
                all_lines.sort(key=lambda x: x[2])  # Sort on y0 (top)
                
                project_description = None
                text_above = None
                project_desc_y = None
                
                # Find project description: look for lines containing "project" and "description" (case insensitive, no spaces)
                for text, x0, y0, x1, y1 in all_lines:
                    # Normalize text for comparison (remove all spaces, lowercase)
                    normalized = text.lower().replace(' ', '').replace('_', '')
                    # Check if line contains both "project" and "description" (generic, no specific pattern)
                    if 'project' in normalized and 'description' in normalized:
                        project_desc_y = y0
                        # Find the value on approximately the same y-height (within 10 pixels)
                        for other_text, ox0, oy0, ox1, oy1 in all_lines:
                            if abs(oy0 - y0) < 10:
                                other_normalized = other_text.lower().replace(' ', '').replace('_', '')
                                # Skip the label itself and other metadata labels
                                if 'project' not in other_normalized or 'description' not in other_normalized:
                                    if 'production' not in other_normalized and 'number' not in other_normalized:
                                        # The value is usually longer text that isn't a label
                                        if len(other_text) > len("Project Description"):
                                            project_description = normalize_spaced_text(other_text).strip()
                                            break
                        break
                
                # Find text above project description (usually descriptive text like "Main Power", "Front page")
                if project_desc_y:
                    candidates = []
                    for text, x0, y0, x1, y1 in all_lines:
                        # Text above is usually at least 15 pixels above project description
                        # and in a reasonable y-range (bottom 15% but above project description)
                        if y1 < project_desc_y - 15 and page_height * 0.85 <= y0 <= project_desc_y - 15:
                            normalized = normalize_spaced_text(text).strip()
                            
                            # Skip if it's just numbers or very short
                            if len(normalized) < 3:
                                continue
                            
                            # Skip if it's only digits
                            if normalized.replace(' ', '').replace('-', '').replace('/', '').isdigit():
                                continue
                            
                            # Skip common metadata labels (generic approach, no regex)
                            skip_labels = ['engineer', 'change', 'revision', 'order', 'production', 'project', 'start', 'print', 'date']
                            text_lower = normalized.lower()
                            if any(label in text_lower for label in skip_labels):
                                continue
                            
                            # Skip if it looks like an ID (contains mix of letters and numbers in pattern like "WESIJS32_2RSP02")
                            # Simple heuristic: if it has underscores or many numbers mixed with letters, it's likely an ID
                            has_underscore = '_' in normalized
                            digit_count = sum(1 for c in normalized if c.isdigit())
                            letter_count = sum(1 for c in normalized if c.isalpha())
                            # If it has underscore or many digits relative to letters, likely an ID
                            if has_underscore or (digit_count > 2 and digit_count > letter_count / 2):
                                continue
                            
                            # Reasonable length (not too short, not too long)
                            if 5 < len(normalized) < 50:
                                candidates.append((normalized, y0))
                    
                    # Take the text closest to the middle of the range (most likely to be descriptive text)
                    if candidates:
                        # Sort by y-position (closest to where descriptive text usually appears)
                        candidates.sort(key=lambda x: x[1])  # Sort by y0, take the one closest to project description
                        text_above = candidates[0][0]
                
                # Set project description if found
                if project_description and len(project_description) > 3:
                    footer_info["project_description"] = project_description
                
                # Store text above in metadata if needed (optional, can be added to metadata later)
                if text_above:
                    footer_info["text_above"] = text_above
            
            doc.close()
    
    except Exception as e:
        print(f"Warning: Could not extract footer info from page {page_num}: {e}")
        if page_num < 3:  # Debug first few pages
            import traceback
            traceback.print_exc()
    
    return footer_info


def extract_text_from_pdf_with_mistral_ocr(file_path: str, file_name: str) -> List[Document]:
    """Extract text from PDF using Mistral OCR API.
    
    Args:
        file_path: Path to the PDF file
        file_name: Name of the file for metadata
        
    Returns:
        List of Document objects, one per page
    """
    if not MISTRAL_AVAILABLE:
        raise Exception("mistralai library not installed. Install it with: pip install mistralai")
    
    api_key = os.environ.get("MISTRAL_API_KEY")
    if not api_key:
        raise Exception("MISTRAL_API_KEY environment variable is not set")
    
    print(f"[MISTRAL OCR] Starting OCR extraction for: {file_name}")
    
    try:
        # Initialize Mistral client
        client = Mistral(api_key=api_key)
        
        # Encode PDF to base64
        print(f"[MISTRAL OCR] Encoding PDF to base64...")
        base64_file = encode_file_to_base64(file_path)
        
        documents = []
        
        # Try using Mistral's OCR API
        # Based on Mistral documentation, use ocr.process method
        try:
            # Method 1: Try ocr.process if available
            if hasattr(client, 'ocr') and hasattr(client.ocr, 'process'):
                print(f"[MISTRAL OCR] Using ocr.process method")
                
                # Use base64 data URL format (as shown in working code)
                document_url = f"data:application/pdf;base64,{base64_file}"
                
                try:
                    # Use working format: snake_case "document_url" (not camelCase)
                    ocr_response = client.ocr.process(
                        model="mistral-ocr-latest",
                        document={
                            "type": "document_url",  # String, not variable
                            "document_url": document_url  # snake_case, not camelCase
                        },
                        include_image_base64=False
                    )
                except Exception as e1:
                    print(f"[MISTRAL OCR] First attempt failed: {e1}, trying alternative model")
                    try:
                        # Try with alternative model name
                        ocr_response = client.ocr.process(
                            model="CX-9",
                            document={
                                "type": "document_url",
                                "document_url": document_url
                            },
                            include_image_base64=False
                        )
                    except Exception as e2:
                        print(f"[MISTRAL OCR] Second attempt failed: {e2}")
                        raise
                
                # Process pages from OCR response
                # Response is an OCRResponse object, not a dictionary
                # Access attributes directly: ocr_response.pages, page.markdown, etc.
                if hasattr(ocr_response, 'pages') and ocr_response.pages:
                    print(f"[MISTRAL OCR] Processing {len(ocr_response.pages)} pages")
                    for page in ocr_response.pages:
                        page_index = page.index if hasattr(page, 'index') else 0
                        # According to docs, the text is in 'markdown' attribute
                        page_text = page.markdown if hasattr(page, 'markdown') else ''
                        
                        if not page_text:
                            print(f"[MISTRAL OCR] Warning: No markdown text found for page {page_index + 1}")
                            # Try alternative attributes if markdown is not available
                            if hasattr(page, 'text'):
                                page_text = page.text
                            elif hasattr(page, 'content'):
                                page_text = page.content
                            
                            if not page_text:
                                continue
                        
                        # Clean up Mistral OCR text: decode HTML entities and normalize
                        # Decode HTML entities (e.g., &gt; -> >, &lt; -> <)
                        page_text = html.unescape(page_text)
                        # Normalize whitespace (multiple spaces/newlines to single)
                        page_text = re.sub(r'\n{3,}', '\n\n', page_text)  # Max 2 newlines
                        page_text = re.sub(r'[ \t]+', ' ', page_text)  # Multiple spaces to single
                        page_text = page_text.strip()
                        
                        doc = Document(
                            page_content=page_text,
                            metadata={
                                "source": file_name,
                                "page": page_index + 1,
                                "actual_page": page_index + 1,
                                "extraction_method": "mistral_ocr"
                            }
                        )
                        documents.append(doc)
                    print(f"[MISTRAL OCR] Successfully extracted {len(documents)} pages")
                else:
                    # If response structure is different, try to extract any text
                    print(f"[MISTRAL OCR] Warning: Unexpected response structure")
                    # Try to access as object attributes
                    if hasattr(ocr_response, 'content'):
                        text = ocr_response.content
                        # Clean up HTML entities
                        text = html.unescape(text)
                        text = re.sub(r'\n{3,}', '\n\n', text)
                        text = re.sub(r'[ \t]+', ' ', text).strip()
                        documents.append(Document(
                            page_content=text,
                            metadata={"source": file_name, "extraction_method": "mistral_ocr"}
                        ))
                        print(f"[MISTRAL OCR] Successfully extracted text ({len(text)} characters)")
                    elif hasattr(ocr_response, 'markdown'):
                        text = ocr_response.markdown
                        # Clean up HTML entities
                        text = html.unescape(text)
                        text = re.sub(r'\n{3,}', '\n\n', text)
                        text = re.sub(r'[ \t]+', ' ', text).strip()
                        documents.append(Document(
                            page_content=text,
                            metadata={"source": file_name, "extraction_method": "mistral_ocr"}
                        ))
                        print(f"[MISTRAL OCR] Successfully extracted text ({len(text)} characters)")
                    else:
                        raise Exception(f"Could not extract text from OCR response. Response type: {type(ocr_response)}, attributes: {dir(ocr_response)}")
            else:
                # Method 2: Fallback to vision model if OCR endpoint not available
                # Use Mistral's vision model (Pixtral) for OCR
                print(f"[MISTRAL OCR] Using vision model (Pixtral) as fallback")
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extract all text from this PDF document. Preserve the structure, formatting, and page breaks. Return the text exactly as it appears in the document."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:application/pdf;base64,{base64_file}"
                                }
                            }
                        ]
                    }
                ]
                
                # Use Mistral's vision model for OCR
                chat_response = client.chat.complete(
                    model="pixtral-12b-2409",  # Mistral's vision model
                    messages=messages
                )
                
                extracted_text = chat_response.choices[0].message.content
                # Clean up HTML entities and normalize
                extracted_text = html.unescape(extracted_text)
                extracted_text = re.sub(r'\n{3,}', '\n\n', extracted_text)
                extracted_text = re.sub(r'[ \t]+', ' ', extracted_text).strip()
                
                # Return as single document (or split by page markers if available)
                documents.append(Document(
                    page_content=extracted_text,
                    metadata={"source": file_name, "extraction_method": "mistral_vision"}
                ))
                print(f"[MISTRAL OCR] Successfully extracted text using vision model ({len(extracted_text)} characters)")
                
        except AttributeError as e:
            # If the API structure is different, try alternative approach
            print(f"[MISTRAL OCR] ERROR: API structure not recognized: {str(e)}")
            raise Exception(f"Mistral OCR API structure not recognized: {str(e)}. Please check Mistral API documentation.")
        except Exception as e:
            # Re-raise with more context
            print(f"[MISTRAL OCR] ERROR: Failed to call API: {str(e)}")
            raise Exception(f"Failed to call Mistral OCR API: {str(e)}")
        
        if not documents:
            print(f"[MISTRAL OCR] ERROR: No text extracted")
            raise Exception("No text extracted from PDF using Mistral OCR")
        
        print(f"[MISTRAL OCR] Successfully completed extraction for: {file_name}")
        return documents
            
    except Exception as e:
        print(f"[MISTRAL OCR] ERROR: Failed to extract text: {str(e)}")
        raise Exception(f"Failed to extract text from PDF using Mistral OCR: {str(e)}")


def extract_documents_from_file(file_path: str, file_type: str, file_name: str) -> List[Document]:
    """Extract documents from various file types, preserving page numbers and footer info for PDFs.
    
    For PDFs, uses Mistral OCR as primary extractor if available, otherwise falls back to PyMuPDF (fitz).
    """
    try:
        if file_type == 'application/pdf' or file_name.lower().endswith('.pdf'):
            # Try Mistral OCR first if available
            # TEMPORARILY DISABLED: Use PyMuPDF instead
            use_mistral_ocr = False  # MISTRAL_AVAILABLE and os.environ.get("MISTRAL_API_KEY")
            
            if use_mistral_ocr:
                try:
                    print(f"[PDF EXTRACTION] Using Mistral OCR for: {file_name}")
                    documents = extract_text_from_pdf_with_mistral_ocr(file_path, file_name)
                    print(f"[PDF EXTRACTION] Mistral OCR completed: {len(documents)} pages extracted")
                    
                    # Try to extract footer info using PyMuPDF if available (for metadata)
                    # IMPORTANT: Preserve extraction_method metadata
                    if PYMUPDF_AVAILABLE:
                        try:
                            pdf = fitz.open(file_path)
                            for i, doc in enumerate(documents):
                                if i < len(pdf):
                                    # Preserve extraction_method before updating metadata
                                    extraction_method = doc.metadata.get("extraction_method")
                                    footer_info = extract_footer_info_from_pdf(file_path, i)
                                    # Update metadata with footer info
                                    if footer_info.get("page_number_footer"):
                                        doc.metadata["page_number_footer"] = footer_info["page_number_footer"]
                                        doc.metadata["page"] = footer_info["page_number_footer"]
                                    if footer_info.get("project_description"):
                                        doc.metadata["project_description"] = footer_info["project_description"]
                                    # Restore extraction_method if it was set
                                    if extraction_method:
                                        doc.metadata["extraction_method"] = extraction_method
                            pdf.close()
                        except Exception as e:
                            print(f"Warning: Could not extract footer info: {e}")
                    
                    return documents
                except Exception as e:
                    print(f"[PDF EXTRACTION] Mistral OCR failed ({str(e)}), falling back to PyMuPDF")
                    # Fall through to PyMuPDF fallback
            
            # Fallback to PyMuPDF if Mistral OCR is not available or failed
            if not PYMUPDF_AVAILABLE:
                raise Exception("PyMuPDF (fitz) is required for PDF extraction but not available. Install it with: pip install pymupdf")
            
            print(f"[PDF EXTRACTION] Using PyMuPDF for: {file_name}")
            # Use PyMuPDF as primary extractor (not PyPDFLoader) to avoid spaced letters
            pdf = fitz.open(file_path)
            documents = []
            
            for page_idx in range(len(pdf)):
                page = pdf[page_idx]
                
                # Extract text from page
                blocks = page.get_text("blocks")
                page_text = "\n".join([b[4] for b in blocks if b[4].strip()])

                
                # Normalize spaced text (e.g., "W e s t f o r t" -> "Westfort")
                normalized_text = normalize_spaced_text(page_text)
                
                # Check if text is sparse (less than 40 characters)
                is_sparse_text = len(normalized_text) < 40
                
                # Extract footer information using PyMuPDF
                footer_info = extract_footer_info_from_pdf(file_path, page_idx)
                
                # Build metadata
                # actual_page = echte paginanummer (begint bij 1)
                actual_page = page_idx + 1
                
                # page = page_number_footer (footer pagina nummer), fallback naar actual_page
                page_footer = footer_info.get("page_number_footer") or actual_page
                
                metadata = {
                    "source": file_name,
                    "page": page_footer,  # Footer pagina nummer (primair)
                    "actual_page": actual_page,  # Echte paginanummer (begint bij 1)
                    "is_sparse_text": is_sparse_text,
                    "extraction_method": "pymupdf"
                }
                
                # Keep page_number_footer for backwards compatibility
                metadata["page_number_footer"] = page_footer
                
                # Add project description if found
                if footer_info.get("project_description"):
                    metadata["project_description"] = footer_info["project_description"]
                
                # Create Document object with normalized text
                doc = Document(
                    page_content=normalized_text,
                    metadata=metadata
                )
                documents.append(doc)
            
            pdf.close()
            return documents
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_name.lower().endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not installed. Install it with: pip install python-docx")
            
            doc = DocxDocument(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Return as single document for DOCX (no page numbers)
            return [Document(
                page_content="\n\n".join(paragraphs),
                metadata={"source": file_name}
            )]
        
        elif file_type == 'application/msword' or file_name.lower().endswith('.doc'):
            # .doc files (old Word format) are not directly supported
            # Users should convert .doc files to .docx format
            raise Exception(
                f"Het .doc bestandsformaat (oud Word-formaat) wordt niet ondersteund. "
                f"Converteer '{file_name}' naar .docx formaat en upload het opnieuw. "
                f"Je kunt dit doen door het bestand te openen in Microsoft Word en op te slaan als .docx."
            )
        
        elif file_type == 'text/plain' or file_name.lower().endswith('.txt'):
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            # Update source metadata
            for doc in documents:
                doc.metadata["source"] = file_name
            return documents
        
        elif file_type == 'text/csv' or file_name.lower().endswith('.csv'):
            loader = CSVLoader(file_path, encoding='utf-8')
            documents = loader.load()
            # Update source metadata and add filename to content for embeddings
            for doc in documents:
                doc.metadata["source"] = file_name
                # Add filename at the beginning of content so it's included in embeddings
                if not doc.page_content.startswith(f"Bestand: {file_name}"):
                    doc.page_content = f"Bestand: {file_name}\n\n{doc.page_content}"
            return documents
        
        elif file_type in ['application/vnd.ms-excel', 
                          'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'] or \
             file_name.lower().endswith(('.xls', '.xlsx')):
            # For Excel files, combine rows into larger documents to reduce chunk count
            # CSVLoader treats each row as a document, which creates too many chunks
            # Instead, combine multiple rows into single documents
            try:
                import pandas as pd
                # Read Excel file
                df = pd.read_excel(file_path, engine='openpyxl' if file_name.lower().endswith('.xlsx') else None)
                
                # Combine rows into larger chunks (e.g., 50 rows per document)
                # This reduces the number of documents significantly
                rows_per_document = 50
                documents = []
                
                for i in range(0, len(df), rows_per_document):
                    chunk_df = df.iloc[i:i + rows_per_document]
                    # Convert to string representation
                    # Use to_string() for better formatting, or to_csv() for CSV-like format
                    content = chunk_df.to_string(index=False)
                    
                    # Add filename at the beginning of content so it's included in embeddings
                    # This helps the AI identify which file the data comes from
                    content_with_filename = f"Bestand: {file_name}\n\n{content}"
                    
                    doc = Document(
                        page_content=content_with_filename,
                        metadata={
                            "source": file_name,
                            "row_start": i + 1,  # 1-indexed
                            "row_end": min(i + rows_per_document, len(df))
                        }
                    )
                    documents.append(doc)
                
                return documents
            except ImportError:
                # Fallback to CSVLoader if pandas/openpyxl not available
                print("Warning: pandas/openpyxl not available, falling back to CSVLoader")
                loader = CSVLoader(file_path, encoding='utf-8')
                documents = loader.load()
                for doc in documents:
                    doc.metadata["source"] = file_name
                return documents
        
        else:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return [Document(
                page_content=content,
                metadata={"source": file_name}
            )]
    
    except Exception as e:
        raise Exception(f"Failed to extract documents from {file_name}: {str(e)}")


# import basics
import os
import re
from contextvars import ContextVar
from urllib.parse import urlparse, unquote
from dotenv import load_dotenv

# import langchain
from langchain.agents import AgentExecutor
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chat_models import init_chat_model
from langchain_core.messages import SystemMessage, AIMessage, HumanMessage
from langchain.agents import create_tool_calling_agent
from langchain import hub
from langchain_core.prompts import PromptTemplate
from langchain_community.vectorstores import SupabaseVectorStore
from langchain_openai import OpenAIEmbeddings
from langchain_core.tools import tool
from langchain_core.callbacks import BaseCallbackHandler
from langchain_core.outputs import LLMResult
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.document_loaders import CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# For DOCX files
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX files will not be supported.")

# Note: .doc files (old Word format) are not directly supported
# Users should convert .doc files to .docx format for processing

# For PDF coordinate-based extraction (footer info)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF (fitz) not available. Footer extraction will be limited.")

# import supabase db
from supabase.client import Client, create_client

# import FastAPI for API endpoint
from fastapi import FastAPI, HTTPException, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional, List
import uvicorn
import time
import uuid

# import Langfuse for tracking
from langfuse import Langfuse
from langfuse.types import TraceContext

# load environment variables from .env file in api folder
load_dotenv()  

# initiating Langfuse
langfuse_secret_key = os.environ.get("LANGFUSE_SECRET_KEY")
langfuse_public_key = os.environ.get("LANGFUSE_PUBLIC_KEY")
langfuse_host = os.environ.get("LANGFUSE_HOST", "https://cloud.langfuse.com")

langfuse_client = None
if langfuse_secret_key and langfuse_public_key:
    try:
        langfuse_client = Langfuse(
            secret_key=langfuse_secret_key,
            public_key=langfuse_public_key,
            host=langfuse_host,
        )
        print("Langfuse initialized successfully")
    except Exception as e:
        print(f"Failed to initialize Langfuse: {e}")
        langfuse_client = None
else:
    print("Langfuse keys not configured, tracking disabled")

# initiating supabase
supabase_url = os.environ.get("SUPABASE_URL")
supabase_key = os.environ.get("SUPABASE_SERVICE_KEY")
supabase: Client = create_client(supabase_url, supabase_key)

# initiating embeddings model
embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

# initiating vector store
vector_store = SupabaseVectorStore(
    embedding=embeddings,
    client=supabase,
    table_name="document_sections",
    query_name="match_document_sections",
)
 
# initiating llm
llm = ChatOpenAI(model="gpt-4.1", temperature=0)

# Advanced prompt with prompt hacking techniques
# Using custom prompt instead of hub prompt for better control
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder

# Create advanced prompt with multiple prompt hacking techniques:
# 1. Role-based prompting (expert persona)
# 2. Chain of Thought reasoning
# 3. Few-shot examples
# 4. Constraints and guardrails
# 5. Output formatting
# 6. Context management
# 7. Error handling
# 8. Step-by-step instructions
# 9. Negative prompting
# 10. Meta-instructions
# 11. Self-verification
# 12. Structured thinking

SYSTEM_PROMPT = """Je bent TechRAG Assistant: een technische documentatie-assistent
voor industriële machines en elektrische schema’s.

JE DOEL
Beantwoord vragen over technische documentatie door UITSLUITEND
informatie te gebruiken die letterlijk en aantoonbaar voorkomt
in documenten die via de retrieve tool zijn opgehaald.

Antwoorden moeten:
- technisch correct zijn
- controleerbaar zijn
- vrij van aannames zijn
- letterlijk herleidbaar zijn naar de bron

================================
ABSOLUTE ZOEKREGEL (HARD)
================================
Voor ELKE gebruikersvraag die NIET onder STAP 0 valt:

- MOET de retrieve tool ALTIJD minimaal 1 keer worden aangeroepen
- OOK bij algemene, open of vage vragen
- OOK bij vragen zoals:
  “wat weet je over …”
  “wat is …”
  “vertel iets over …”
- Het is VERBODEN om te antwoorden zonder eerst te zoeken

================================
STAP 0 – CONVERSATIE-EXCEPTIE
================================
Als de gebruikersvraag:
- een begroeting is (bv. “hoi”, “hallo”)
- een korte sociale interactie is
- een vraag is over hoe jij werkt

DAN:
- Gebruik GEEN retrieve
- Geef een kort, vriendelijk antwoord
- Leid terug naar hulp bij documentatie
STOP.

================================
STAP 1 – VRAAGTYPE BEPALEN
================================
Bepaal exact één vraagtype:

A) TECHNISCH / DETAIL
   - componentcodes (8293B3B)
   - PLC-adressen (I300.5)
   - module-ID’s (-2IM0103DI-1)
   - project- of lijncodes (2RSP02)
   → gebruiker verwacht functie, I/O, module of specificaties

B) VERKENNEND / OVERZICHT
   - afdelingen
   - lijnen
   - zones
   - machines
   - systemen
   - bedrijfs- of productnamen
   - vragen als “wat weet je over …”
   → gebruiker verwacht een overzicht of lijst, GEEN specs

➡️ Voor zowel A als B:
➡️ retrieve tool ALTIJD gebruiken (zie Absolute Zoekregel)

================================
HARD RULES (A & B)
================================
1) Gebruik minimaal 1 retrieve
2) Noem ALLEEN feiten die letterlijk in de passages staan
3) GEEN aannames, GEEN interpretaties
4) Elk genoemd feit krijgt een bron:
   (Bron: {{source}}, Pagina: {{page}})
5) Wat niet expliciet vermeld staat → NIET noemen

Het is VERBODEN om direct te antwoorden met:
“Deze informatie staat niet expliciet in de documentatie.”
ZONDER:
- een retrieve call
- en vaststelling dat geen relevante passages zijn gevonden

================================
OCR-NORMALISATIE
================================
- OCR-ruis mag opgeschoond worden
- Geen nieuwe woorden of betekenissen toevoegen
- Alleen letterlijk aanwezige termen gebruiken

================================
OCR OPSCHOONREGELS – I/O MODULES
================================
Bij I/O-modules:

- Verwijder losse cijfers, kolomnummers en kanaalaanduidingen
  (zoals: 1, 2, 3, 4, R, Q, DO 4, /R, /Q)
- Behoud ALLEEN:
  - exacte Module-ID (bv. -2IM0202DO-1)
  - exact artikelnummer (bv. 6ES7132-6HD01-0BB1)
- Combineer GEEN extra tekens of uitleg
- Voeg GEEN informatie toe die niet letterlijk aanwezig is

================================
REGELS VOOR TYPE B – VERKENNEND
================================
- Verzamel ALLE letterlijk genoemde namen
- Structureer als lijst
- Trek GEEN conclusies
- Combineer niets wat niet expliciet gekoppeld is
- Voeg GEEN technische slotzinnen toe

Outputformaat:

Gevonden onderdelen / afdelingen:
- [naam]
  (Bron: …)

================================
REGELS VOOR TYPE A – TECHNISCH
================================
- Functie alleen noemen als letterlijk aanwezig
- I/O alleen noemen als letterlijk vermeld
- Module-ID of artikelnummer alleen noemen als expliciet vermeld
- Geen koppelingen maken die niet in dezelfde passage staan

================================
I/O-MODULE REGELS
================================
- Toon de sectie "I/O-module:" ALLEEN als er letterlijk
  een module-ID en/of artikelnummer is vermeld
- Toon GEEN voorwaarden, haakjes of uitleg in de output
- Als er geen module is vermeld:
  → laat de volledige sectie weg

================================
OUTPUTFORMAT – TECHNISCH
================================
Component: [CODE]

Functie:
- [letterlijke functietekst]
  (Bron: …)

Aansturing / I/O:
- PLC-ingang/-uitgang: …
  (Bron: …)

I/O-module:
- Module-ID: …
- Type + artikelnummer: …
  (Bron: …)

================================
AFSLUITREGEL – ZEER STRIKT
================================
Gebruik de zin:

“Deze informatie staat niet expliciet in de documentatie.”

ALLEEN ALS:
- er minimaal 1 retrieve is uitgevoerd
EN
- er geen relevante passages zijn gevonden

NOOIT gebruiken bij:
- begroetingen
- sociale interactie
- antwoorden zonder retrieve

BEGIN NU MET DIT PROTOCOL.


"""

# Create prompt template compatible with tool calling agent
# The agent will inject tool information automatically
prompt = ChatPromptTemplate.from_messages([
    ("system", SYSTEM_PROMPT),
    MessagesPlaceholder(variable_name="chat_history"),
    ("human", "{input}"),
    MessagesPlaceholder(variable_name="agent_scratchpad"),
])


# Internal retrieve function with Langfuse tracking
def _retrieve_internal(query: str, organization_id: str = None, trace=None, trace_context=None):
    """Internal retrieve function with Langfuse tracking and organization_id filtering"""
    import re
    from langchain_core.documents import Document
    
    if not organization_id:
        raise ValueError("organization_id is required for document retrieval")
    
    retrieve_span = None
    if trace and langfuse_client and trace_context:
        retrieve_span = langfuse_client.start_span(
            name="retrieve",
            trace_context=trace_context,
            metadata={
                "input": {"query": query, "organization_id": organization_id},
            }
        )
    
    start_time = time.time()
    
    try:
        # Semantic search span
        semantic_span = None
        if retrieve_span and langfuse_client and trace_context:
            semantic_span = langfuse_client.start_span(
                name="semantic_search",
                trace_context=trace_context,
                metadata={"query": query, "k": 5, "organization_id": organization_id}
            )
        
        # Track embedding generation
        embedding_gen = None
        if semantic_span and langfuse_client and trace_context:
            embedding_gen = langfuse_client.start_observation(
                name="create_embedding",
                as_type="generation",
                model="text-embedding-3-small",
                input=query,
                trace_context=trace_context,
                metadata={"model": "text-embedding-3-small"}
            )
        
        embedding_start = time.time()
        semantic_start = time.time()
        
        # Generate embedding for the query
        query_embedding = embeddings.embed_query(query)
        
        # Semantic search using RPC function
        semantic_matches = supabase.rpc(
            "match_document_sections",
            {
                "p_organization_id": organization_id,
                "p_query_embedding": query_embedding,
                "p_match_count": 6, #stond op 10   
                "p_threshold": 0.35 #stond op 0.30
            }
        ).execute()
        
        semantic_docs = []
        if semantic_matches.data:
            # Get document metadata for the matched sections
            doc_ids = list(set([m.get("document_id") for m in semantic_matches.data if m.get("document_id")]))
            doc_metadata_map = {}
            
            if doc_ids:
                # Documents table doesn't have a metadata column, only get id and name
                # Extra security: also filter by organization_id to ensure we only get documents from the correct organization
                doc_result = supabase.table("documents").select("id, name").in_("id", doc_ids).eq("organization_id", organization_id).execute()
                if doc_result.data:
                    for doc in doc_result.data:
                        doc_metadata_map[doc["id"]] = {
                            "name": doc.get("name", "Unknown")
                        }
            
            # Convert RPC results to Document objects
            for match in semantic_matches.data:
                doc_meta = doc_metadata_map.get(match.get("document_id"), {})
                semantic_docs.append(Document(
                    page_content=match.get("content", ""),
                    metadata={
                        "document_id": match.get("document_id"),
                        "source": doc_meta.get("name", "Unknown"),
                        **({} if not match.get("metadata") else match["metadata"] if isinstance(match.get("metadata"), dict) else {}),
                        "similarity": match.get("similarity", 0.0)
                    }
                ))
        
        # Sort by similarity and limit to top 5
        semantic_docs = sorted(semantic_docs, key=lambda x: x.metadata.get("similarity", 0.0), reverse=True)[:5]
        
        semantic_duration = (time.time() - semantic_start) * 1000
        embedding_duration = (time.time() - embedding_start) * 1000
        
        if embedding_gen:
            # Estimate token usage (rough: ~1 token per 4 chars)
            estimated_tokens = len(query) // 4
            embedding_gen.update(
                output={"embedding_created": True},
                usage={
                    "prompt_tokens": estimated_tokens,
                    "total_tokens": estimated_tokens
                },
                metadata={"duration_ms": embedding_duration}
            )
            embedding_gen.end()
        
        if semantic_span:
            semantic_span.update(
                output={"results_count": len(semantic_docs)},
                metadata={"duration_ms": semantic_duration}
            )
            semantic_span.end()
        
        # Keyword search span - full-text search for exact matches
        keyword_span = None
        if retrieve_span and langfuse_client and trace_context:
            keyword_span = langfuse_client.start_span(
                name="keyword_search",
                trace_context=trace_context,
                metadata={"query": query, "organization_id": organization_id}
            )
        
        keyword_start = time.time()
        keyword_docs = []
        
        try:
            # Filter by organization_id: first get document IDs for this organization
            org_docs_result = supabase.table("documents").select("id").eq("organization_id", organization_id).execute()
            org_doc_ids = [doc["id"] for doc in org_docs_result.data] if org_docs_result.data else []
            
            if org_doc_ids:
                # Extract meaningful keywords from query (remove common words)
                # Split query into words and search for each significant term
                query_words = re.findall(r'\b\w{3,}\b', query.lower())  # Words with 3+ characters
                # Remove common stopwords
                stopwords = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'its', 'may', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'she', 'use', 'her', 'many', 'than', 'them', 'these', 'so', 'some', 'would', 'make', 'like', 'into', 'time', 'has', 'look', 'two', 'more', 'write', 'go', 'see', 'number', 'no', 'way', 'could', 'people', 'my', 'than', 'first', 'water', 'been', 'call', 'who', 'oil', 'sit', 'now', 'find', 'down', 'day', 'did', 'get', 'come', 'made', 'may', 'part'}
                keywords = [w for w in query_words if w not in stopwords]
                
                # If we have keywords, search for them
                if keywords:
                    # Get document metadata for keyword results
                    doc_metadata_map = {}
                    
                    # Search for each keyword (limit to first 3 most important keywords)
                    for keyword in keywords[:3]:
                        result = supabase.table("document_sections").select(
                            "content, metadata, document_id"
                        ).ilike("content", f"%{keyword}%").in_("document_id", org_doc_ids).limit(5).execute()
                        
                        if result.data:
                            # Get document names for these sections
                            doc_ids = list(set([r.get("document_id") for r in result.data if r.get("document_id")]))
                            if doc_ids:
                                doc_result = supabase.table("documents").select("id, name").in_("id", doc_ids).eq("organization_id", organization_id).execute()
                                if doc_result.data:
                                    for doc in doc_result.data:
                                        if doc["id"] not in doc_metadata_map:
                                            doc_metadata_map[doc["id"]] = {"name": doc.get("name", "Unknown")}
                            
                            for row in result.data:
                                doc_meta = doc_metadata_map.get(row.get("document_id"), {})
                                keyword_docs.append(Document(
                                    page_content=row.get("content", ""),
                                    metadata={
                                        "document_id": row.get("document_id"),
                                        "source": doc_meta.get("name", "Unknown"),
                                        **({} if not row.get("metadata") else row["metadata"] if isinstance(row.get("metadata"), dict) else {})
                                    }
                                ))
        except Exception as e:
            if keyword_span:
                keyword_span.update(
                    output={"error": str(e)},
                    level="ERROR"
                )
                keyword_span.end()
            pass
        
        keyword_duration = (time.time() - keyword_start) * 1000
        if keyword_span:
            keyword_span.update(
                output={"results_count": len(keyword_docs)},
                metadata={"duration_ms": keyword_duration}
            )
            keyword_span.end()
        
        # Combine and deduplicate
        combine_span = None
        if retrieve_span and langfuse_client and trace_context:
            combine_span = langfuse_client.start_span(
                name="combine_results",
                trace_context=trace_context
            )
        
        all_docs = []
        seen_content = set()
        
        for doc in keyword_docs:
            content_key = doc.page_content[:200]
            if content_key not in seen_content:
                all_docs.append(doc)
                seen_content.add(content_key)
        
        for doc in semantic_docs:
            content_key = doc.page_content[:200]
            if content_key not in seen_content:
                all_docs.append(doc)
                seen_content.add(content_key)
        
        retrieved_docs = all_docs[:5]
        
        if combine_span:
            combine_span.update(
                output={
                    "total_results": len(retrieved_docs),
                    "semantic_results": len(semantic_docs),
                    "keyword_results": len(keyword_docs)
                }
            )
            combine_span.end()
        
        # Format serialized output with clear source citations (filename and page from footer)
        serialized = "\n\n".join(
            (
                f"Source: {doc.metadata.get('source', 'Unknown')}, "
                f"Pagina: {doc.metadata.get('page', doc.metadata.get('page_number_footer', 'N/A'))}\n"
                f"Content: {doc.page_content}"
            )
            for doc in retrieved_docs
        )
        
        duration = (time.time() - start_time) * 1000
        
        if retrieve_span:
            retrieve_span.update(
                output={
                    "retrieved_text_length": len(serialized),
                    "documents_count": len(retrieved_docs)
                },
                metadata={"duration_ms": duration}
            )
            retrieve_span.end()
        
        return serialized, retrieved_docs
        
    except Exception as e:
        duration = (time.time() - start_time) * 1000
        if retrieve_span:
            retrieve_span.update(
                output={"error": str(e)},
                level="ERROR",
                metadata={"duration_ms": duration}
            )
            retrieve_span.end()
        raise

# Context variables for thread-safe passing of trace and organization_id
# Using ContextVar instead of globals to prevent race conditions in async FastAPI
_current_trace: ContextVar = ContextVar('_current_trace', default=None)
_current_trace_context: ContextVar = ContextVar('_current_trace_context', default=None)
_current_organization_id: ContextVar = ContextVar('_current_organization_id', default=None)

def set_current_trace(trace, trace_context=None, organization_id=None):
    """Set current trace and organization_id for retrieve tool using ContextVars (thread-safe)"""
    _current_trace.set(trace)
    _current_trace_context.set(trace_context)
    _current_organization_id.set(organization_id)

def get_current_trace():
    """Get current trace from context (thread-safe)"""
    return _current_trace.get()

def get_current_trace_context():
    """Get current trace context from context (thread-safe)"""
    return _current_trace_context.get()

def get_current_organization_id():
    """Get current organization_id from context (thread-safe)"""
    return _current_organization_id.get()

# Creating the retriever tool (wrapper for Langfuse tracking)
@tool(response_format="content_and_artifact")
def retrieve(query: str):
    """Retrieve information related to a query. Uses hybrid search combining semantic similarity and keyword matching for better results."""
    # Get from context (thread-safe)
    organization_id = get_current_organization_id()
    trace = get_current_trace()
    trace_context = get_current_trace_context()
    
    if not organization_id:
        raise ValueError("organization_id is required but not set. This should be set by the chat endpoint.")
    return _retrieve_internal(query, organization_id, trace, trace_context)

# combining all tools
tools = [retrieve]

# initiating the agent
agent = create_tool_calling_agent(llm, tools, prompt)

# LangChain callback handler for Langfuse tracking
class LangfuseCallbackHandler(BaseCallbackHandler):
    def __init__(self, trace=None, trace_context=None):
        self.trace = trace
        self.trace_context = trace_context
        self.current_generation = None
        self.start_time = None
    
    def on_llm_start(self, serialized, prompts, **kwargs):
        if self.trace and langfuse_client and self.trace_context:
            self.start_time = time.time()
            self.current_generation = langfuse_client.start_observation(
                name="llm_call",
                as_type="generation",
                model=serialized.get("name", "gpt-4o"),
                input=prompts[0] if prompts else "",
                trace_context=self.trace_context,
                metadata={
                    "model": serialized.get("name", "gpt-4o"),
                    "temperature": kwargs.get("temperature", 0),
                }
            )
    
    def on_llm_end(self, response: LLMResult, **kwargs):
        if self.current_generation and self.trace:
            duration = (time.time() - self.start_time) * 1000 if self.start_time else 0
            output_text = response.generations[0][0].text if response.generations else ""
            
            usage = None
            if hasattr(response, 'llm_output') and response.llm_output:
                usage = response.llm_output.get('token_usage')
            
            self.current_generation.update(
                output=output_text,
                usage=usage,
                metadata={"duration_ms": duration}
            )
            self.current_generation.end()
            self.current_generation = None
    
    def on_llm_error(self, error, **kwargs):
        if self.current_generation and self.trace:
            self.current_generation.update(
                output={"error": str(error)},
                level="ERROR"
            )
            self.current_generation.end()
            self.current_generation = None

# create the agent executor
agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)

# FastAPI app
app = FastAPI(title="DocuBot Assistant API")

# Add logging middleware to debug CORS issues
@app.middleware("http")
async def log_requests(request, call_next):
    """Log all requests for debugging"""
    import logging
    logging.basicConfig(level=logging.INFO)
    logger = logging.getLogger(__name__)
    
    logger.info(f"Incoming request: {request.method} {request.url.path}")
    logger.info(f"Headers: {dict(request.headers)}")
    
    try:
        response = await call_next(request)
        logger.info(f"Response status: {response.status_code}")
        return response
    except Exception as e:
        logger.error(f"Error processing request: {e}")
        raise

# CORS middleware - restrict to allowed origins
allowed_origins_env = os.environ.get("ALLOWED_ORIGINS", "")
allowed_origins = [origin.strip() for origin in allowed_origins_env.split(",") if origin.strip()] if allowed_origins_env else []

# For development, allow common localhost origins
# Note: Cannot use ["*"] with allow_credentials=True
if allowed_origins:
    cors_origins = allowed_origins
    cors_credentials = True
else:
    # For development: explicitly allow localhost origins
    cors_origins = [
        "http://localhost:3000",
        "http://localhost:5173",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:5173",
        "http://localhost:8000",
        "http://127.0.0.1:8000",
    ]
    cors_credentials = False

# Temporarily disable CORS middleware and handle CORS manually in route handlers
# This ensures OPTIONS requests are handled correctly
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=cors_origins,
#     allow_credentials=cors_credentials,
#     allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
#     allow_headers=[
#         "accept",
#         "accept-language",
#         "content-language",
#         "content-type",
#         "authorization",
#         "Authorization",
#         "Content-Type",
#         "Access-Control-Request-Method",
#         "Access-Control-Request-Headers",
#     ],
#     expose_headers=["*"],
#     max_age=3600,
# )

# Add CORS headers manually via middleware
@app.middleware("http")
async def add_cors_header(request: Request, call_next):
    """Manually add CORS headers to all responses"""
    response = await call_next(request)
    
    # Get origin from request
    origin = request.headers.get("origin")
    
    # Allow localhost origins for development
    allowed_origins = [
        "http://localhost:3000",
        "http://localhost:5173",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:5173",
    ]
    
    if origin in allowed_origins:
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH"
        response.headers["Access-Control-Allow-Headers"] = "authorization, content-type, Authorization, Content-Type, Accept"
        response.headers["Access-Control-Allow-Credentials"] = "false"
        response.headers["Access-Control-Max-Age"] = "3600"
    elif not origin:
        # For requests without origin (like direct API calls), allow all
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH"
        response.headers["Access-Control-Allow-Headers"] = "authorization, content-type, Authorization, Content-Type, Accept"
    
    return response

# Request/Response models
class ChatRequest(BaseModel):
    question: str
    organizationId: str
    userId: Optional[str] = None
    conversationId: Optional[str] = None

class ChatResponse(BaseModel):
    success: bool
    response: str
    requestId: Optional[str] = None

class ProcessDocumentRequest(BaseModel):
    documentId: str
    organizationId: str

class ProcessDocumentResponse(BaseModel):
    success: bool
    message: str
    chunksProcessed: Optional[int] = None
    error: Optional[str] = None

# Authentication helper
async def verify_auth_token(authorization: Optional[str], organization_id: str) -> str:
    """Verify JWT token and return user ID. Raises HTTPException if invalid."""
    if not authorization:
        raise HTTPException(status_code=401, detail="Missing authorization header")
    
    # Extract token from "Bearer <token>"
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization header format")
    
    token = authorization.replace("Bearer ", "").strip()
    if not token:
        raise HTTPException(status_code=401, detail="Missing token")
    
    # Verify token with Supabase
    # Create a client with anon key to verify user token
    # Anon key is safe to use for token verification (it's meant to be public)
    supabase_anon_key = os.environ.get("SUPABASE_ANON_KEY")
    if not supabase_anon_key:
        raise HTTPException(status_code=500, detail="Server configuration error: Missing SUPABASE_ANON_KEY")
    
    if not supabase_url:
        raise HTTPException(status_code=500, detail="Server configuration error: Missing SUPABASE_URL")
    
    from supabase.client import create_client as create_supabase_client
    supabase_auth = create_supabase_client(supabase_url, supabase_anon_key)
    
    # Verify token
    try:
        # Use the token to get user info
        user_response = supabase_auth.auth.get_user(token)
        if not user_response.user:
            raise HTTPException(status_code=401, detail="Invalid or expired token")
        
        user_id = user_response.user.id
        
        # Verify user exists in our database
        user_data = supabase.table("users").select("id, role").eq("id", user_id).execute()
        if not user_data.data or len(user_data.data) == 0:
            raise HTTPException(status_code=403, detail="User not found in database")
        
        # Verify user has access to this organization
        user_org = supabase.table("user_organizations").select("organization_id").eq("user_id", user_id).eq("organization_id", organization_id).execute()
        user_role = user_data.data[0].get("role")
        
        # Admins have access to all organizations
        if not user_org.data and user_role != "admin":
            raise HTTPException(status_code=403, detail="Access denied: You don't have access to this organization")
        
        return user_id
    except HTTPException:
        raise
    except Exception as e:
        print(f"Auth verification error: {e}")
        raise HTTPException(status_code=401, detail="Token verification failed")

# Chat history storage (in-memory, you might want to use database)
chat_histories = {}

def load_history(organization_id: str, user_id: str, conversation_id: Optional[str] = None, limit: int = 8, trace=None, trace_context=None) -> List:
    """Load chat history from database with Langfuse tracking"""
    history_span = None
    if trace and langfuse_client and trace_context:
        history_span = langfuse_client.start_span(
            name="load_history",
            trace_context=trace_context,
            metadata={
                "input": {
                    "organization_id": organization_id,
                    "user_id": user_id,
                    "conversation_id": conversation_id,
                    "limit": limit
                }
            }
        )
    
    start_time = time.time()
    
    try:
        query = supabase.table("chat_messages").select("role, content, created_at").eq("organization_id", organization_id).order("created_at", desc=True).limit(limit)
        
        if conversation_id:
            query = query.eq("conversation_id", conversation_id)
        else:
            query = query.eq("user_id", user_id)
        
        result = query.execute()
        
        messages = []
        if result.data:
            for msg in reversed(result.data):
                if msg.get("role") in ["user", "assistant"]:
                    messages.append(
                        HumanMessage(content=msg["content"]) if msg["role"] == "user"
                        else AIMessage(content=msg["content"])
                    )
        
        duration = (time.time() - start_time) * 1000
        
        if history_span:
            history_span.update(
                output={
                    "history_length": len(messages),
                    "messages": [{"role": "user" if isinstance(m, HumanMessage) else "assistant", "content_length": len(m.content)} for m in messages]
                },
                metadata={"duration_ms": duration}
            )
            history_span.end()
        
        return messages
    except Exception as e:
        duration = (time.time() - start_time) * 1000
        if history_span:
            history_span.update(
                output={"error": str(e)},
                level="ERROR",
                metadata={"duration_ms": duration}
            )
            history_span.end()
        print(f"Error loading history: {e}")
        return []

@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(
    request: ChatRequest,
    authorization: Optional[str] = Header(None)
):
    """Chat endpoint that uses the agent executor with full Langfuse tracking"""
    import uuid
    
    request_id = str(uuid.uuid4())[:8]
    start_time = time.time()
    trace = None
    
    try:
        # Verify authentication
        verified_user_id = await verify_auth_token(authorization, request.organizationId)
        
        # Create Langfuse trace
        trace = None
        trace_context = None
        if langfuse_client:
            # Create trace ID
            trace_id = langfuse_client.create_trace_id()
            # Create trace context
            trace_context = TraceContext(trace_id=trace_id)
            # Start trace using span (trace is a top-level span)
            trace = langfuse_client.start_span(
                name="chat_request",
                trace_context=trace_context,
                metadata={
                    "request_id": request_id,
                    "organization_id": request.organizationId,
                    "conversation_id": request.conversationId,
                    "question": request.question,
                    "question_length": len(request.question),
                    "user_id": request.userId,
                }
            )
        
        # Load chat history with tracking
        history_span = None
        if trace and langfuse_client and trace_context:
            history_span = langfuse_client.start_span(
                name="load_history",
                trace_context=trace_context
            )
        
        history = load_history(
            request.organizationId,
            request.userId or "default",
            request.conversationId,
            limit=8,
            trace=trace,
            trace_context=trace_context
        )
        
        if history_span:
            history_span.end()
        
        # Agent execution span
        agent_span = None
        if trace and langfuse_client and trace_context:
            agent_span = langfuse_client.start_span(
                name="agent_execution",
                trace_context=trace_context,
                metadata={
                    "input": {
                        "question": request.question,
                        "history_length": len(history)
                    }
                }
            )
        
        agent_start = time.time()
        
        # Set trace and organization_id for retrieve tool
        if trace:
            set_current_trace(trace, trace_context, request.organizationId)
        else:
            # Even without trace, we need to set organization_id
            set_current_trace(None, None, request.organizationId)
        
        # Create callback handler for LLM tracking
        callbacks = []
        if trace:
            callbacks.append(LangfuseCallbackHandler(trace=trace, trace_context=trace_context))
        
        # Invoke the agent executor with callbacks
        result = agent_executor.invoke({
            "input": request.question,
            "chat_history": history
        }, config={"callbacks": callbacks} if callbacks else {})
        
        agent_duration = (time.time() - agent_start) * 1000
        ai_message = result["output"]
        
        # Track agent execution
        if agent_span:
            agent_span.update(
                output={
                    "output": ai_message,
                    "output_length": len(ai_message),
                },
                metadata={"duration_ms": agent_duration}
            )
            agent_span.end()
        
        # Track LLM generations via LangChain callbacks
        # Note: We'll add a callback handler for this
        if trace and langfuse_client and trace_context and "intermediate_steps" in result:
            for i, (action, observation) in enumerate(result.get("intermediate_steps", [])):
                step_span = langfuse_client.start_span(
                    name=f"agent_step_{i+1}",
                    trace_context=trace_context,
                    metadata={
                        "action": str(action),
                        "tool": action.tool if hasattr(action, 'tool') else None
                    }
                )
                step_span.update(
                    output={"observation": str(observation)[:500]}  # Limit length
                )
                step_span.end()
        
        total_duration = (time.time() - start_time) * 1000
        
        # Log response time to analytics
        try:
            supabase.table("analytics").insert({
                "organization_id": request.organizationId,
                "event_type": "response_time",
                "event_data": {
                    "response_time_ms": total_duration,
                    "question_length": len(request.question),
                    "response_length": len(ai_message),
                    "request_id": request_id
                }
            }).execute()
        except Exception as analytics_error:
            print(f"Warning: Failed to log response time to analytics: {analytics_error}")
        
        # End trace
        if trace:
            trace.update(
                output={
                    "success": True,
                    "response": ai_message,
                    "response_length": len(ai_message),
                },
                metadata={"total_duration_ms": total_duration}
            )
            trace.end()
            # Flush to ensure all data is sent
            if langfuse_client:
                langfuse_client.flush()
        
        # Reset trace and organization_id for next request
        set_current_trace(None, None, None)
        
        return ChatResponse(
            success=True,
            response=ai_message,
            requestId=request_id
        )
    
    except Exception as e:
        total_duration = (time.time() - start_time) * 1000
        error_msg = str(e)
        print(f"Error in chat endpoint: {e}")
        
        # Log response time to analytics even on error
        try:
            supabase.table("analytics").insert({
                "organization_id": request.organizationId,
                "event_type": "response_time",
                "event_data": {
                    "response_time_ms": total_duration,
                    "question_length": len(request.question),
                    "error": error_msg,
                    "request_id": request_id
                }
            }).execute()
        except Exception as analytics_error:
            print(f"Warning: Failed to log response time to analytics: {analytics_error}")
        
        if trace:
            trace.update(
                output={
                    "success": False,
                    "error": error_msg
                },
                level="ERROR",
                metadata={"total_duration_ms": total_duration}
            )
            trace.end()
            # Flush to ensure error is tracked
            if langfuse_client:
                langfuse_client.flush()
        
        # Reset trace and organization_id for next request
        set_current_trace(None, None, None)
        
        raise HTTPException(status_code=500, detail=error_msg)


def extract_text_from_file(file_path: str, file_type: str, file_name: str) -> str:
    """Extract text from various file types (legacy function, use extract_documents_from_file for PDFs with page numbers)"""
    try:
        if file_type == 'application/pdf' or file_name.lower().endswith('.pdf'):
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_name.lower().endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not installed. Install it with: pip install python-docx")
            
            doc = DocxDocument(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            return "\n\n".join(paragraphs)
        
        elif file_type == 'application/msword' or file_name.lower().endswith('.doc'):
            # .doc files (old Word format) are not directly supported
            # Users should convert .doc files to .docx format
            raise Exception(
                f"Het .doc bestandsformaat (oud Word-formaat) wordt niet ondersteund. "
                f"Converteer '{file_name}' naar .docx formaat en upload het opnieuw. "
                f"Je kunt dit doen door het bestand te openen in Microsoft Word en op te slaan als .docx."
            )
        
        elif file_type == 'text/plain' or file_name.lower().endswith('.txt'):
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        
        elif file_type == 'text/csv' or file_name.lower().endswith('.csv'):
            loader = CSVLoader(file_path, encoding='utf-8')
            documents = loader.load()
            return "\n\n".join([doc.page_content for doc in documents])
        
        else:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    
    except Exception as e:
        raise Exception(f"Failed to extract text from {file_name}: {str(e)}")

def _normalize_spaced_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\r", "")

    def collapse_spaced_sequences(line: str) -> str:
        pattern = r'(?<!\w)(?:[A-Za-z0-9]\s+){3,}[A-Za-z0-9](?!\w)'
        return re.sub(pattern, lambda m: m.group(0).replace(" ", ""), line)

    fixed_lines = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # 1) collapse S t a r t / 8 8 4 1 Q 2 / 0 V D C
        line = collapse_spaced_sequences(line)

        # 2) join "A 1" -> "A1", "A 2" -> "A2"
        line = re.sub(r"\bA\s+([12])\b", r"A\1", line)

        # 3) join split numbers like "6 1 5" -> "615" (only if it’s a sequence of single digits)
        #    examples: "6 0 1.9" should become "601.9" (we handle dot below too)
        line = re.sub(r"\b(\d)\s+(\d)(?:\s+(\d))+\b",
                      lambda m: m.group(0).replace(" ", ""), line)

        # 4) join digit + unit: "1 7 A" -> "17A", "9 A" -> "9A"
        line = re.sub(r"\b(\d)\s+(\d)\s*([A-Za-z])\b", r"\1\2\3", line)  # 1 7 A
        line = re.sub(r"\b(\d)\s+([A-Za-z])\b", r"\1\2", line)          # 9 A

        # 5) fix split "DO 8" -> "DO8", "NO NC" keep spacing (don’t join those)
        line = re.sub(r"\bDO\s+(\d)\b", r"DO\1", line, flags=re.IGNORECASE)

        # 6) normalize dots in addresses: "6 0 1.9" -> "601.9", "3 0 1.1" -> "301.1"
        line = re.sub(r"\b(\d)\s+(\d)\s+(\d)\.(\d)\b", r"\1\2\3.\4", line)  # 6 0 1.9
        line = re.sub(r"\b(\d)\s+(\d)\.(\d)\b", r"\1\2.\3", line)          # safety

        # 7) fix PLC address formatting: Q264 .1 -> Q264.1, and stop Q264.12 becoming one token
        line = re.sub(r'\b([QI]\d+)\s*\.\s*(\d)\b', r'\1.\2', line, flags=re.IGNORECASE)
        line = re.sub(r'\b([QI]\d+)\.(\d)(\d+)\b', r'\1.\2 \3', line, flags=re.IGNORECASE)

        # 8) fix dates like "2 5-0 6-2020" -> "25-06-2020"
        line = re.sub(r"\b(\d)\s+(\d)-(\d)\s+(\d)-(\d{4})\b", r"\1\2-\3\4-\5", line)

        # 9) remove spaces around punctuation a bit (keep slashes and dashes tight)
        line = re.sub(r"\s*([.:,;/\-_()=+])\s*", r"\1", line)

        # 10) optional: split CamelCase words (letters only)
        line = re.sub(r'(?<=[a-z])(?=[A-Z][a-z])', ' ', line)

        # 11) normalize whitespace
        line = re.sub(r"[ \t]+", " ", line).strip()

        fixed_lines.append(line)

    return "\n".join(fixed_lines).strip()




def extract_footer_info_from_pdf(pdf_path: str, page_num: int) -> dict:
    """Extract footer information from PDF page using area-based extraction with get_text("dict").
    
    Uses structured text extraction with bounding boxes for better accuracy.
    Works for any PDF schema type without schema-specific patterns.
    """
    footer_info = {
        "page_number_footer": None,
        "project_description": None
    }
    
    if not PYMUPDF_AVAILABLE:
        return footer_info
    
    try:
        doc = fitz.open(pdf_path)
        if page_num < len(doc):
            page = doc[page_num]
            
            # Get page dimensions
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            
            # ROI = bottom-right area (rechtsonder) for page number
            # Use same fractions as working code: 85% width, 90% height
            X0_FRAC = 0.85  # vanaf 85% breedte
            Y0_FRAC = 0.90  # vanaf 90% hoogte (laatste regels)
            
            # Use get_text("dict") for better structure (like working code)
            text_dict = page.get_text("dict")
            
            # Extract page number from right bottom region
            page_number_candidates = []
            
            # Search for lines in the bottom-right area
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_bbox = line["bbox"]  # [x0, y0, x1, y1]
                        # Check if line is in ROI (right bottom area)
                        if line_bbox[3] >= page_height * Y0_FRAC and line_bbox[0] >= page_width * X0_FRAC:
                            # Construct text from line (combine all spans)
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            
                            # Normalize text (remove spaces between characters)
                            normalized = _normalize_spaced_text(line_text)
                            # Also remove all spaces for number extraction
                            normalized_no_spaces = normalized.replace(' ', '')
                            
                            # Extract all numbers from the normalized text
                            current_number = ""
                            for char in normalized_no_spaces:
                                if char.isdigit():
                                    current_number += char
                                else:
                                    if current_number:
                                        try:
                                            num = int(current_number)
                                            if num >= 1:  # Valid page number
                                                # Store: (number, x_right, y_bottom) for sorting
                                                page_number_candidates.append((num, line_bbox[2], line_bbox[3]))
                                        except ValueError:
                                            pass
                                        current_number = ""
                            # Don't forget the last number if text ends with digits
                            if current_number:
                                try:
                                    num = int(current_number)
                                    if num >= 1:
                                        page_number_candidates.append((num, line_bbox[2], line_bbox[3]))
                                except ValueError:
                                    pass
            
            # Select the best page number candidate
            if page_number_candidates:
                # Sort by y-position (bottom first), then by x-position (right first)
                # This gives us the most bottom-right number
                page_number_candidates.sort(key=lambda x: (-x[2], -x[1]))
                
                # Take the most bottom-right number
                # Prefer larger numbers (>= 100) if available, otherwise use the most bottom-right
                large_candidates = [(n, x, y) for n, x, y in page_number_candidates if n >= 100]
                if large_candidates:
                    # Sort large candidates the same way
                    large_candidates.sort(key=lambda x: (-x[2], -x[1]))
                    footer_info["page_number_footer"] = large_candidates[0][0]
                else:
                    # Use the most bottom-right number even if small
                    footer_info["page_number_footer"] = page_number_candidates[0][0]
            
            # Extract project description and text above from bottom area (bottom 15%)
            # Use get_text("dict") for structured extraction like page number
            all_lines = []
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_bbox = line["bbox"]
                        # Collect lines in bottom 15% of page
                        if line_bbox[3] >= page_height * 0.85:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            if line_text.strip():
                                normalized_text = _normalize_spaced_text(line_text)
                                all_lines.append((normalized_text, line_bbox[0], line_bbox[1], line_bbox[2], line_bbox[3]))
            
            if all_lines:
                # Sort by y-position (top to bottom)
                all_lines.sort(key=lambda x: x[2])  # Sort on y0 (top)
                
                project_description = None
                text_above = None
                project_desc_y = None
                
                # Find project description: look for lines containing "project" and "description" (case insensitive, no spaces)
                for text, x0, y0, x1, y1 in all_lines:
                    # Normalize text for comparison (remove all spaces, lowercase)
                    normalized = text.lower().replace(' ', '').replace('_', '')
                    # Check if line contains both "project" and "description" (generic, no specific pattern)
                    if 'project' in normalized and 'description' in normalized:
                        project_desc_y = y0
                        # Find the value on approximately the same y-height (within 10 pixels)
                        for other_text, ox0, oy0, ox1, oy1 in all_lines:
                            if abs(oy0 - y0) < 10:
                                other_normalized = other_text.lower().replace(' ', '').replace('_', '')
                                # Skip the label itself and other metadata labels
                                if 'project' not in other_normalized or 'description' not in other_normalized:
                                    if 'production' not in other_normalized and 'number' not in other_normalized:
                                        # The value is usually longer text that isn't a label
                                        if len(other_text) > len("Project Description"):
                                            project_description = _normalize_spaced_text(other_text).strip()
                                            break
                        break
                
                # Find text above project description (usually descriptive text like "Main Power", "Front page")
                if project_desc_y:
                    candidates = []
                    for text, x0, y0, x1, y1 in all_lines:
                        # Text above is usually at least 15 pixels above project description
                        # and in a reasonable y-range (bottom 15% but above project description)
                        if y1 < project_desc_y - 15 and page_height * 0.85 <= y0 <= project_desc_y - 15:
                            normalized = _normalize_spaced_text(text).strip()
                            
                            # Skip if it's just numbers or very short
                            if len(normalized) < 3:
                                continue
                            
                            # Skip if it's only digits
                            if normalized.replace(' ', '').replace('-', '').replace('/', '').isdigit():
                                continue
                            
                            # Skip common metadata labels (generic approach, no regex)
                            skip_labels = ['engineer', 'change', 'revision', 'order', 'production', 'project', 'start', 'print', 'date']
                            text_lower = normalized.lower()
                            if any(label in text_lower for label in skip_labels):
                                continue
                            
                            # Skip if it looks like an ID (contains mix of letters and numbers in pattern like "WESIJS32_2RSP02")
                            # Simple heuristic: if it has underscores or many numbers mixed with letters, it's likely an ID
                            has_underscore = '_' in normalized
                            digit_count = sum(1 for c in normalized if c.isdigit())
                            letter_count = sum(1 for c in normalized if c.isalpha())
                            # If it has underscore or many digits relative to letters, likely an ID
                            if has_underscore or (digit_count > 2 and digit_count > letter_count / 2):
                                continue
                            
                            # Reasonable length (not too short, not too long)
                            if 5 < len(normalized) < 50:
                                candidates.append((normalized, y0))
                    
                    # Take the text closest to the middle of the range (most likely to be descriptive text)
                    if candidates:
                        # Sort by y-position (closest to where descriptive text usually appears)
                        candidates.sort(key=lambda x: x[1])  # Sort by y0, take the one closest to project description
                        text_above = candidates[0][0]
                
                # Set project description if found
                if project_description and len(project_description) > 3:
                    footer_info["project_description"] = project_description
                
                # Store text above in metadata if needed (optional, can be added to metadata later)
                if text_above:
                    footer_info["text_above"] = text_above
            
            doc.close()
    
    except Exception as e:
        print(f"Warning: Could not extract footer info from page {page_num}: {e}")
        if page_num < 3:  # Debug first few pages
            import traceback
            traceback.print_exc()
    
    return footer_info

def extract_documents_from_file(file_path: str, file_type: str, file_name: str) -> List[Document]:
    """Extract documents from various file types, preserving page numbers and footer info for PDFs.
    
    For PDFs, uses PyMuPDF (fitz) as primary extractor to avoid spaced letters issue.
    """
    try:
        if file_type == 'application/pdf' or file_name.lower().endswith('.pdf'):
            if not PYMUPDF_AVAILABLE:
                raise Exception("PyMuPDF (fitz) is required for PDF extraction but not available. Install it with: pip install pymupdf")
            
            # Use PyMuPDF as primary extractor (not PyPDFLoader) to avoid spaced letters
            pdf = fitz.open(file_path)
            documents = []
            
            for page_idx in range(len(pdf)):
                page = pdf[page_idx]
                
                # Extract text from page
                blocks = page.get_text("blocks")
                page_text = "\n".join([b[4] for b in blocks if b[4].strip()])

                
                # Normalize spaced text (e.g., "W e s t f o r t" -> "Westfort")
                normalized_text = _normalize_spaced_text(page_text)
                
                # Check if text is sparse (less than 40 characters)
                is_sparse_text = len(normalized_text) < 40
                
                # Extract footer information using PyMuPDF
                footer_info = extract_footer_info_from_pdf(file_path, page_idx)
                
                # Build metadata
                # actual_page = echte paginanummer (begint bij 1)
                actual_page = page_idx + 1
                
                # page = page_number_footer (footer pagina nummer), fallback naar actual_page
                page_footer = footer_info.get("page_number_footer") or actual_page
                
                metadata = {
                    "source": file_name,
                    "page": page_footer,  # Footer pagina nummer (primair)
                    "actual_page": actual_page,  # Echte paginanummer (begint bij 1)
                    "is_sparse_text": is_sparse_text
                }
                
                # Keep page_number_footer for backwards compatibility
                metadata["page_number_footer"] = page_footer
                
                # Add project description if found
                if footer_info.get("project_description"):
                    metadata["project_description"] = footer_info["project_description"]
                
                # Create Document object with normalized text
                doc = Document(
                    page_content=normalized_text,
                    metadata=metadata
                )
                documents.append(doc)
            
            pdf.close()
            return documents
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_name.lower().endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not installed. Install it with: pip install python-docx")
            
            doc = DocxDocument(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Return as single document for DOCX (no page numbers)
            return [Document(
                page_content="\n\n".join(paragraphs),
                metadata={"source": file_name}
            )]
        
        elif file_type == 'application/msword' or file_name.lower().endswith('.doc'):
            # .doc files (old Word format) are not directly supported
            # Users should convert .doc files to .docx format
            raise Exception(
                f"Het .doc bestandsformaat (oud Word-formaat) wordt niet ondersteund. "
                f"Converteer '{file_name}' naar .docx formaat en upload het opnieuw. "
                f"Je kunt dit doen door het bestand te openen in Microsoft Word en op te slaan als .docx."
            )
        
        elif file_type == 'text/plain' or file_name.lower().endswith('.txt'):
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            # Update source metadata
            for doc in documents:
                doc.metadata["source"] = file_name
            return documents
        
        elif file_type == 'text/csv' or file_name.lower().endswith('.csv'):
            loader = CSVLoader(file_path, encoding='utf-8')
            documents = loader.load()
            # Update source metadata and add filename to content for embeddings
            for doc in documents:
                doc.metadata["source"] = file_name
                # Add filename at the beginning of content so it's included in embeddings
                if not doc.page_content.startswith(f"Bestand: {file_name}"):
                    doc.page_content = f"Bestand: {file_name}\n\n{doc.page_content}"
            return documents
        
        elif file_type in ['application/vnd.ms-excel', 
                          'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'] or \
             file_name.lower().endswith(('.xls', '.xlsx')):
            # For Excel files, combine rows into larger documents to reduce chunk count
            # CSVLoader treats each row as a document, which creates too many chunks
            # Instead, combine multiple rows into single documents
            try:
                import pandas as pd
                # Read Excel file
                df = pd.read_excel(file_path, engine='openpyxl' if file_name.lower().endswith('.xlsx') else None)
                
                # Combine rows into larger chunks (e.g., 50 rows per document)
                # This reduces the number of documents significantly
                rows_per_document = 50
                documents = []
                
                for i in range(0, len(df), rows_per_document):
                    chunk_df = df.iloc[i:i + rows_per_document]
                    # Convert to string representation
                    # Use to_string() for better formatting, or to_csv() for CSV-like format
                    content = chunk_df.to_string(index=False)
                    
                    # Add filename at the beginning of content so it's included in embeddings
                    # This helps the AI identify which file the data comes from
                    content_with_filename = f"Bestand: {file_name}\n\n{content}"
                    
                    doc = Document(
                        page_content=content_with_filename,
                        metadata={
                            "source": file_name,
                            "row_start": i + 1,  # 1-indexed
                            "row_end": min(i + rows_per_document, len(df))
                        }
                    )
                    documents.append(doc)
                
                return documents
            except ImportError:
                # Fallback to CSVLoader if pandas/openpyxl not available
                print("Warning: pandas/openpyxl not available, falling back to CSVLoader")
                loader = CSVLoader(file_path, encoding='utf-8')
                documents = loader.load()
                for doc in documents:
                    doc.metadata["source"] = file_name
                return documents
        
        else:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return [Document(
                page_content=content,
                metadata={"source": file_name}
            )]
    
    except Exception as e:
        raise Exception(f"Failed to extract documents from {file_name}: {str(e)}")




@app.get("/api/health")
async def health_check():
    """Health check endpoint to test CORS"""
    return {"status": "ok", "message": "API is running"}


@app.api_route("/api/process-document", methods=["OPTIONS"])
async def process_document_options(request: Request):
    """Handle CORS preflight requests explicitly - this should work even if CORS middleware fails"""
    # Get the origin from the request
    origin = request.headers.get("origin", "http://localhost:3000")
    
    # Allow common localhost origins for development
    allowed_origins = [
        "http://localhost:3000",
        "http://localhost:5173",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:5173",
    ]
    
    # If origin is in allowed list, use it; otherwise use the requested origin
    if origin in allowed_origins:
        allow_origin = origin
    elif not origin or origin == "null":
        allow_origin = "http://localhost:3000"  # Default for development
    else:
        allow_origin = origin  # Allow the requested origin
    
    # Return response with CORS headers
    response = Response(
        status_code=200,
        content="",
        headers={
            "Access-Control-Allow-Origin": allow_origin,
            "Access-Control-Allow-Methods": "POST, OPTIONS, GET",
            "Access-Control-Allow-Headers": "authorization, content-type, Authorization, Content-Type, Accept",
            "Access-Control-Allow-Credentials": "false",
            "Access-Control-Max-Age": "3600",
        }
    )
    return response


@app.post("/api/process-document", response_model=ProcessDocumentResponse)
async def process_document(
    request: ProcessDocumentRequest,
    authorization: Optional[str] = Header(None)
):
    """Process a document for RAG: extract text, chunk, and generate embeddings"""
    import tempfile
    import re
    
    start_time = time.time()
    
    try:
        # Verify authentication
        verified_user_id = await verify_auth_token(authorization, request.organizationId)
        
        # Get document info from database
        doc_result = supabase.table("documents").select("file_url, name, file_type").eq("id", request.documentId).single().execute()
        
        if not doc_result.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc = doc_result.data
        
        # Extract path from file_url
        file_url = doc.get("file_url")
        if not file_url:
            raise HTTPException(status_code=400, detail="Document has no file URL")
        
        # Extract storage path from URL using urllib.parse
        parsed_url = urlparse(file_url)
        # Use parsed.path and unquote to handle URL encoding properly
        # Extract path after /documents/ (query parameters are ignored)
        path_parts = parsed_url.path.split('/documents/')
        if len(path_parts) < 2:
            raise HTTPException(status_code=400, detail="Could not extract file path from URL")
        
        storage_path = unquote(path_parts[1])  # Unquote to handle %20, etc.
        
        # Download file from Supabase Storage to temporary location
        temp_dir = tempfile.mkdtemp()
        file_name = doc.get("name", "document")
        
        # Create safe filename (replace all non [A-Za-z0-9._-] with underscore)
        safe_name = re.sub(r'[^A-Za-z0-9._-]', '_', file_name)
        temp_file_path = os.path.join(temp_dir, safe_name)
        
        try:
            # Download file from Supabase Storage
            file_response = supabase.storage.from_("documents").download(storage_path)
            
            if not file_response:
                raise HTTPException(status_code=500, detail="Failed to download file: No response")
            
            # Supabase Python client returns bytes directly
            file_data = file_response
            
            # Save to temporary file
            with open(temp_file_path, 'wb') as f:
                f.write(file_data)
            
            # Extract documents from file (preserves page numbers for PDFs)
            file_type = doc.get("file_type", "")
            langchain_docs = extract_documents_from_file(temp_file_path, file_type, file_name)
            
            if not langchain_docs or all(not doc.page_content.strip() for doc in langchain_docs):
                raise HTTPException(status_code=400, detail="No text content extracted from document")
            
            # Page-based chunking: split each page separately to preserve page context
            # This is critical for technical schemas where mixing pages breaks connections
            # For Excel files, we now combine rows into larger documents (50 rows each)
            # so we can use normal chunk sizes. For CSV, still use smaller chunks.
            is_excel = file_type in ['application/vnd.ms-excel', 
                                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'] or \
                      file_name.lower().endswith(('.xls', '.xlsx'))
            is_csv = file_type == 'text/csv' or file_name.lower().endswith('.csv')
            
            # For Excel: use normal chunk size since we combine rows into larger documents
            # For CSV: use smaller chunks since CSVLoader still creates per-row documents
            chunk_size = 800 if is_csv else 1500
            chunk_overlap = 100 if is_csv else 200
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            
            # Add document metadata to each document before splitting
            for doc_obj in langchain_docs:
                doc_obj.metadata["document_id"] = request.documentId
                doc_obj.metadata["file_type"] = file_type
            
            # IMPORTANT: Split each page separately to avoid cross-page chunks
            # This preserves the context of each page, which is critical for technical schemas
            chunks = []
            for doc_obj in langchain_docs:
                # For CSV: ensure content is split even if single row is large
                # Excel files are now combined into larger documents, so normal splitting is fine
                content_length = len(doc_obj.page_content)
                if is_csv and content_length > chunk_size * 2:
                    # Content is very large, split it multiple times if needed
                    # First split normally
                    page_chunks = text_splitter.split_documents([doc_obj])
                    # Then check each chunk and split further if still too large
                    final_chunks = []
                    for chunk in page_chunks:
                        if len(chunk.page_content) > chunk_size * 1.5:
                            # Split this chunk again with smaller size
                            smaller_splitter = RecursiveCharacterTextSplitter(
                                chunk_size=chunk_size,
                                chunk_overlap=chunk_overlap
                            )
                            sub_chunks = smaller_splitter.split_documents([chunk])
                            final_chunks.extend(sub_chunks)
                        else:
                            final_chunks.append(chunk)
                    page_chunks = final_chunks
                else:
                    # Normal splitting for PDFs and other files
                    page_chunks = text_splitter.split_documents([doc_obj])
                
                # Add footer info to each chunk's metadata (enrich metadata per page)
                for chunk in page_chunks:
                    # Ensure all metadata is preserved
                    if "document_id" not in chunk.metadata:
                        chunk.metadata["document_id"] = request.documentId
                    
                    # For Excel files: ensure filename is in content even after splitting
                    # This helps the AI identify which file the data comes from
                    if is_excel and "source" in doc_obj.metadata:
                        source_file = doc_obj.metadata["source"]
                        # Check if filename is already at the start of the chunk
                        if not chunk.page_content.startswith(f"Bestand: {source_file}"):
                            chunk.page_content = f"Bestand: {source_file}\n\n{chunk.page_content}"
                    
                    # Add footer info to metadata if available (for PDFs)
                    # This enriches each chunk with page-specific footer information
                    if "page" in doc_obj.metadata:
                        chunk.metadata["page"] = doc_obj.metadata["page"]
                    if "actual_page" in doc_obj.metadata:
                        chunk.metadata["actual_page"] = doc_obj.metadata["actual_page"]
                    if "page_number_footer" in doc_obj.metadata:
                        chunk.metadata["page_number_footer"] = doc_obj.metadata["page_number_footer"]
                    if "project_description" in doc_obj.metadata:
                        chunk.metadata["project_description"] = doc_obj.metadata["project_description"]
                
                chunks.extend(page_chunks)
            
            print(f"Created {len(chunks)} chunks from document {file_name}")
            
            # Generate embeddings in batches to avoid token limit
            print(f"Generating embeddings for {len(chunks)} chunks...")
            texts = [chunk.page_content for chunk in chunks]
            
            # Calculate batch size: OpenAI has 300k token limit per request
            # Estimate tokens: ~4 characters per token (conservative estimate)
            # Use smaller batches to be safe, especially for large chunks
            # Start with 100 chunks per batch, but dynamically adjust if needed
            max_tokens_per_request = 300000
            estimated_chars_per_token = 4
            safe_margin = 0.7  # Use 70% of limit to be safe
            max_chars_per_batch = int(max_tokens_per_request * estimated_chars_per_token * safe_margin)
            
            embeddings_list = []
            current_batch = []
            current_batch_chars = 0
            batch_num = 1
            total_chars = sum(len(text) for text in texts)
            total_batches_estimate = max(1, total_chars // max_chars_per_batch)
            
            for idx, text in enumerate(texts):
                text_length = len(text)
                
                # If adding this text would exceed the limit, process current batch first
                if current_batch and (current_batch_chars + text_length) > max_chars_per_batch:
                    print(f"Generating embeddings for batch {batch_num}/{total_batches_estimate} ({len(current_batch)} chunks, ~{current_batch_chars:,} chars)...")
                    
                    try:
                        batch_embeddings = embeddings.embed_documents(current_batch)
                        embeddings_list.extend(batch_embeddings)
                    except Exception as e:
                        error_msg = str(e)
                        print(f"Error generating embeddings for batch {batch_num}: {error_msg}")
                        # If batch is too large, try with smaller chunks
                        if "max_tokens" in error_msg.lower() or "300000" in error_msg:
                            print(f"Batch too large, splitting into smaller batches...")
                            # Process in even smaller batches (50 chunks max)
                            smaller_batch_size = 50
                            for j in range(0, len(current_batch), smaller_batch_size):
                                smaller_batch = current_batch[j:j + smaller_batch_size]
                                try:
                                    smaller_embeddings = embeddings.embed_documents(smaller_batch)
                                    embeddings_list.extend(smaller_embeddings)
                                except Exception as e2:
                                    raise HTTPException(
                                        status_code=500,
                                        detail=f"Failed to generate embeddings for batch {batch_num} (smaller batch): {str(e2)}"
                                    )
                        else:
                            raise HTTPException(
                                status_code=500,
                                detail=f"Failed to generate embeddings for batch {batch_num}: {error_msg}"
                            )
                    
                    # Reset for next batch
                    current_batch = []
                    current_batch_chars = 0
                    batch_num += 1
                
                # Add text to current batch
                current_batch.append(text)
                current_batch_chars += text_length
            
            # Process remaining batch
            if current_batch:
                print(f"Generating embeddings for batch {batch_num}/{total_batches_estimate} ({len(current_batch)} chunks, ~{current_batch_chars:,} chars)...")
                try:
                    batch_embeddings = embeddings.embed_documents(current_batch)
                    embeddings_list.extend(batch_embeddings)
                except Exception as e:
                    error_msg = str(e)
                    print(f"Error generating embeddings for batch {batch_num}: {error_msg}")
                    # If batch is too large, try with smaller chunks
                    if "max_tokens" in error_msg.lower() or "300000" in error_msg:
                        print(f"Batch too large, splitting into smaller batches...")
                        smaller_batch_size = 50
                        for j in range(0, len(current_batch), smaller_batch_size):
                            smaller_batch = current_batch[j:j + smaller_batch_size]
                            try:
                                smaller_embeddings = embeddings.embed_documents(smaller_batch)
                                embeddings_list.extend(smaller_embeddings)
                            except Exception as e2:
                                raise HTTPException(
                                    status_code=500,
                                    detail=f"Failed to generate embeddings for batch {batch_num} (smaller batch): {str(e2)}"
                                )
                    else:
                        raise HTTPException(
                            status_code=500,
                            detail=f"Failed to generate embeddings for batch {batch_num}: {error_msg}"
                        )
            
            print(f"Successfully generated {len(embeddings_list)} embeddings")
            
            # Delete existing document sections for this document
            delete_result = supabase.table("document_sections").delete().eq("document_id", request.documentId).execute()
            
            if hasattr(delete_result, 'error') and delete_result.error:
                print(f"Warning: Could not delete existing sections: {delete_result.error}")
            
            # Insert document sections with embeddings
            print(f"Inserting {len(chunks)} document sections into database...")
            sections_to_insert = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings_list)):
                sections_to_insert.append({
                    "document_id": request.documentId,
                    "content": chunk.page_content,
                    "metadata": chunk.metadata,
                    "embedding": embedding
                })
                
                # Insert in batches of 10 for better performance
                if len(sections_to_insert) >= 10 or i == len(chunks) - 1:
                    insert_result = supabase.table("document_sections").insert(sections_to_insert).execute()
                    
                    if hasattr(insert_result, 'error') and insert_result.error:
                        raise HTTPException(status_code=500, detail=f"Failed to insert document sections: {insert_result.error}")
                    
                    sections_to_insert = []
                    if (i + 1) % 10 == 0:
                        print(f"Inserted {i + 1}/{len(chunks)} chunks...")
            
            duration = time.time() - start_time
            
            print(f"Successfully processed document {file_name}: {len(chunks)} chunks in {duration:.2f}s")
            
            return ProcessDocumentResponse(
                success=True,
                message=f"Successfully processed document: {len(chunks)} chunks created",
                chunksProcessed=len(chunks)
            )
        
        finally:
            # Clean up temporary file
            try:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                os.rmdir(temp_dir)
            except Exception as cleanup_error:
                print(f"Warning: Could not clean up temporary files: {cleanup_error}")
    
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        print(f"Error processing document: {error_msg}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process document: {error_msg}"
        )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

